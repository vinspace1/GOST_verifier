from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass, asdict
from typing import Any, Dict, List, Optional, Tuple, Iterator

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn


# -------------------- Поиск изображений в run --------------------
def extract_image_rids_from_run(run) -> List[str]:
    r = run._r
    blips = r.xpath(".//*[local-name()='blip']")
    rids: List[str] = []
    for blip in blips:
        rid = blip.get(qn("r:embed"))
        if rid:
            rids.append(rid)
    return rids


def resolve_image_part(doc: Document, rid: str) -> Tuple[bytes, str]:
    part = doc.part.related_parts[rid]  # ImagePart
    blob = part.blob
    filename = str(part.partname).split("/")[-1]  # /word/media/image1.png -> image1.png
    return blob, filename


# -------------------- Подписи (caption) --------------------
CAPTION_STYLE_NAMES = {
    "caption",
    "подпись",
    "подпись рисунка",
    "подпись таблицы",
}

FIG_CAPTION_RE = re.compile(
    r"^\s*(?:Рисунок|Рис\.|Fig\.|Figure)\s*(?P<num>\d+(?:\.\d+)*)\s*[-—:.)]?\s*(?P<title>.+?)\s*$",
    re.IGNORECASE
)

def is_caption_style(p: Paragraph) -> bool:
    name = (p.style.name or "").strip().lower()
    if name in CAPTION_STYLE_NAMES:
        return True
    return ("caption" in name) or ("подпись" in name)

def parse_figure_caption(text: str) -> Optional[Dict[str, str]]:
    m = FIG_CAPTION_RE.match(text or "")
    if not m:
        return None
    return {"number": m.group("num"), "title": m.group("title").strip()}


# -------------------- Важно: обходим ВСЕ абзацы (в т.ч. внутри таблиц) --------------------
def iter_paragraphs_in_table(table: Table) -> Iterator[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            # абзацы в ячейке
            for p in cell.paragraphs:
                yield p
            # вложенные таблицы
            for t in cell.tables:
                yield from iter_paragraphs_in_table(t)

def iter_all_paragraphs(doc: Document) -> Iterator[Paragraph]:
    # абзацы верхнего уровня
    for p in doc.paragraphs:
        yield p
    # абзацы внутри таблиц верхнего уровня
    for t in doc.tables:
        yield from iter_paragraphs_in_table(t)

def paragraph_has_image(p: Paragraph) -> bool:
    for run in p.runs:
        if extract_image_rids_from_run(run):
            return True
    return False


# Чтобы найти подпись "рядом", нам нужен список всех Paragraph в порядке обхода
def build_paragraph_stream(doc: Document) -> List[Paragraph]:
    return list(iter_all_paragraphs(doc))

def pick_caption_nearby(
    paras: List[Paragraph],
    idx: int,
    prefer_below: bool = True,
    window: int = 2,
) -> Tuple[str, str, str]:
    """
    Ищем подпись рядом с абзацем idx (который содержит картинку).
    Сначала вниз, потом вверх (или наоборот).
    """
    def try_para(i: int) -> Optional[Tuple[str, str, str]]:
        if i < 0 or i >= len(paras):
            return None
        p = paras[i]
        txt = (p.text or "").strip()
        if not txt:
            return None

        if is_caption_style(p):
            parsed = parse_figure_caption(txt)
            if parsed:
                return txt, parsed["number"], parsed["title"]
            return txt, "", ""

        parsed = parse_figure_caption(txt)
        if parsed:
            return txt, parsed["number"], parsed["title"]

        return None

    order: List[int] = []
    for k in range(1, window + 1):
        if prefer_below:
            order += [idx + k, idx - k]
        else:
            order += [idx - k, idx + k]

    for j in order:
        res = try_para(j)
        if res:
            return res

    return "", "", ""


# -------------------- Результат --------------------
@dataclass
class FoundImage:
    image_index: int
    rid: str
    paragraph_index: int  # индекс в потоке всех paragraph (включая таблицы)
    run_index: int
    saved_path: str
    original_filename: str
    caption: str = ""
    caption_number: str = ""
    caption_title: str = ""


def extract_images_to_folder_and_json(
    docx_path: str,
    out_dir: str = "images_out",
    json_path: str = "images.json",
    prefer_caption_below: bool = True,
    caption_window: int = 2,
) -> Dict[str, Any]:
    os.makedirs(out_dir, exist_ok=True)

    doc = Document(docx_path)
    paras = build_paragraph_stream(doc)

    images: List[FoundImage] = []
    image_counter = 0

    for pi, p in enumerate(paras):
        # быстрый пропуск, если нет картинок
        if not paragraph_has_image(p):
            continue

        for run_idx, run in enumerate(p.runs):
            rids = extract_image_rids_from_run(run)
            for rid in rids:
                image_counter += 1
                blob, orig_name = resolve_image_part(doc, rid)

                safe_name = f"img_{image_counter:04d}_{orig_name}"
                saved_path = os.path.join(out_dir, safe_name)
                with open(saved_path, "wb") as f:
                    f.write(blob)

                cap, cap_num, cap_title = pick_caption_nearby(
                    paras, idx=pi,
                    prefer_below=prefer_caption_below,
                    window=caption_window,
                )

                images.append(FoundImage(
                    image_index=image_counter,
                    rid=rid,
                    paragraph_index=pi,
                    run_index=run_idx,
                    saved_path=saved_path,
                    original_filename=orig_name,
                    caption=cap,
                    caption_number=cap_num,
                    caption_title=cap_title,
                ))

    payload = {
        "source_docx": docx_path,
        "output_dir": out_dir,
        "images_total": len(images),
        "images": [asdict(x) for x in images],
    }

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    return payload


if __name__ == "__main__":
    result = extract_images_to_folder_and_json(
        "original.docx",       # <-- твой файл
        out_dir="images_out",
        json_path="images.json",
        prefer_caption_below=True,
        caption_window=2,
    )
    print(f"OK: extracted {result['images_total']} images -> {result['output_dir']}, json -> images.json")
